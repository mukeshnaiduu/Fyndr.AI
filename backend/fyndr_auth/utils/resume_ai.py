import io
import json
import logging
import os
import re
from typing import Any, Dict, List, Optional, Tuple

from django.conf import settings

logger = logging.getLogger(__name__)


def _read_pdf_text(data: bytes) -> str:
	try:
		from pypdf import PdfReader
		reader = PdfReader(io.BytesIO(data))
		texts: List[str] = []
		for page in reader.pages:
			try:
				t = page.extract_text() or ""
			except Exception:
				t = ""
			if t:
				texts.append(t)
		return "\n".join(texts).strip()
	except Exception as e:
		logger.warning(f"PDF text extraction failed: {e}")
		return ""


def _read_docx_text(data: bytes) -> str:
	try:
		# python-docx requires a file-like object pointing to a path; BytesIO works
		from docx import Document
		bio = io.BytesIO(data)
		doc = Document(bio)
		paras = [p.text for p in doc.paragraphs if p.text]
		return "\n".join(paras).strip()
	except Exception as e:
		logger.warning(f"DOCX text extraction failed: {e}")
		return ""


def _find_urls_in_text(text: str) -> List[str]:
	# Simple regex to find http(s), www, and mailto links
	if not text:
		return []
	regex = re.compile(r"(https?://[^\s)\]\"'<>]+|www\.[^\s)\]\"'<>]+|mailto:[^\s)\]\"'<>]+)", re.IGNORECASE)
	matches = regex.findall(text)
	cleaned = set()
	for m in matches:
		u = m.strip().rstrip('.,;:')
		if u:
			# prepend http:// for www. links to normalize
			if u.lower().startswith('www.'):
				u = 'http://' + u
			cleaned.add(u)
	return list(cleaned)


def _extract_links_from_pdf(data: bytes, text: str = '') -> List[str]:
	links = set()
	try:
		from pypdf import PdfReader
		reader = PdfReader(io.BytesIO(data))
		for page in reader.pages:
			try:
				annots = page.get('/Annots')
			except Exception:
				annots = None
			if annots:
				for a in annots:
					try:
						obj = a.get_object()
						# /A -> action dictionary, /URI key holds link
						action = obj.get('/A') if isinstance(obj, dict) else None
						if action and isinstance(action, dict) and action.get('/URI'):
							links.add(action.get('/URI'))
					except Exception:
						try:
							# fallback: some annots provide /URI directly
							uri = a.get('/URI')
							if uri:
								links.add(uri)
						except Exception:
							continue
	except Exception:
		# ignore PDF link extraction errors
		pass

	# Also scan visible text for URLs
	try:
		for u in _find_urls_in_text(text):
			links.add(u)
	except Exception:
		pass

	# Fallback: search raw PDF bytes for URI occurrences if nothing found yet
	try:
		if not links:
			bs = data
			try:
				# quick bytes to str decode preserving ascii
				s = bs.decode('latin1')
			except Exception:
				s = ''
			# reuse existing url finder to avoid regex quoting issues
			for u in _find_urls_in_text(s):
				links.add(u)
	except Exception:
		pass

	return list(links)


def _extract_links_from_docx(data: bytes, text: str = '') -> List[str]:
	links = set()
	try:
		from docx import Document
		bio = io.BytesIO(data)
		doc = Document(bio)
		# relationships on document part often include hyperlinks
		try:
			rels = doc.part.rels
			for r in rels.values():
				if r.reltype and 'hyperlink' in r.reltype:
					target = getattr(r, 'target_ref', None) or getattr(r, 'target', None)
					if target:
						links.add(target)
		except Exception:
			pass
	except Exception:
		# ignore docx link extraction errors
		pass

	# fallback: scan text for urls
	try:
		for u in _find_urls_in_text(text):
			links.add(u)
	except Exception:
		pass

	return list(links)


def extract_text_from_resume(content_type: str, filename: str, data: bytes) -> str:
	"""Backward-compatible: return only text (keeps existing callers working)."""
	ct = (content_type or "").lower()
	name = (filename or "").lower()
	if ct.endswith("pdf") or name.endswith(".pdf"):
		return _read_pdf_text(data)
	# Handle legacy .doc before checking for generic 'word' content-types
	if name.endswith(".doc"):
		# Legacy .doc: if it's a true OLE file, we can't parse without antiword; if not OLE (fake .doc), try text decode.
		is_ole = data.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1")
		if is_ole:
			logger.info(".doc (OLE) detected; skipping text extraction without antiword")
			return ""
		# Non-OLE .doc (often plain text saved with .doc extension). Attempt safe decode.
		try:
			text = data.decode("utf-8", errors="ignore")
			if text and any(ch.isalpha() for ch in text):
				return text
		except Exception:
			pass
		try:
			text = data.decode("latin-1", errors="ignore")
			return text
		except Exception:
			return ""
	# DOCX (Office Open XML)
	if "vnd.openxmlformats" in ct or name.endswith(".docx"):
		return _read_docx_text(data)
	# Fallback: try decode as text
	try:
		return data.decode("utf-8", errors="ignore")
	except Exception:
		return ""


def extract_text_and_links_from_resume(content_type: str, filename: str, data: bytes) -> Tuple[str, List[str]]:
	"""Return (text, links) extracted from resume bytes. Links collected from annotations/relationships and by regex over text."""
	ct = (content_type or "").lower()
	name = (filename or "").lower()
	text = ""
	links: List[str] = []
	if ct.endswith("pdf") or name.endswith(".pdf"):
		text = _read_pdf_text(data)
		links = _extract_links_from_pdf(data, text)
		return text, links
	# Handle legacy .doc before generic 'word' content-type
	if name.endswith(".doc"):
		is_ole = data.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1")
		if is_ole:
			logger.info(".doc (OLE) detected; skipping text+link extraction without antiword")
			return "", []
		# Non-OLE .doc: treat as plain text and extract URLs by regex
		try:
			text = data.decode("utf-8", errors="ignore")
		except Exception:
			try:
				text = data.decode("latin-1", errors="ignore")
			except Exception:
				text = ""
		try:
			links = _find_urls_in_text(text)
		except Exception:
			links = []
		return text, links
	if "vnd.openxmlformats" in ct or name.endswith(".docx"):
		text = _read_docx_text(data)
		links = _extract_links_from_docx(data, text)
		return text, links
	try:
		text = data.decode("utf-8", errors="ignore")
		links = _find_urls_in_text(text)
		return text, links
	except Exception:
		return "", []


def _heuristic_extract_projects_from_text(text: str) -> List[Dict[str, Any]]:
	"""Try to find a Projects section and extract simple project entries.

	This is a best-effort parser: it looks for a section header like 'Projects' and
	then captures bullet/line-based entries that contain a title and optionally a URL
	and tech tokens. Returns list of {title, link, description, tech_stack, domain}.
	"""
	if not text or not isinstance(text, str):
		return []
	projects: List[Dict[str, Any]] = []
	# Find Projects section (case-insensitive)
	m = re.search(r"(^|\n)\s*(projects|personal projects|selected projects)\s*[:\n]", text, flags=re.IGNORECASE)
	if not m:
		return []
	start = m.end()
	# Extract the following 800-2000 chars as section body
	section = text[start:start+2000]
	# Split lines and look for candidate entries
	lines = [l.strip() for l in section.splitlines() if l.strip()]
	for line in lines:
		# stop if next section header likely appears
		if re.match(r"^[A-Z][A-Za-z ]{1,30}:$", line):
			break
		# Look for URL
		urls = _find_urls_in_text(line)
		link = urls[0] if urls else ''
		# Attempt to split title and description by ' - ' or ' — ' or ':'
		parts = re.split(r"\s[-–—:]\s", line, maxsplit=1)
		title = parts[0]
		description = parts[1] if len(parts) > 1 else ''
		# Remove trailing URL from title/description
		title = re.sub(r"https?://\S+|www\.\S+", "", title).strip().strip('•-–—')
		description = re.sub(r"https?://\S+|www\.\S+", "", description).strip()
		# Extract tech tokens by looking for parentheses or 'Tech:' or common separators
		tech_stack = []
		tech_match = re.search(r"\(([^)]+)\)", line)
		if tech_match:
			tech_stack = [t.strip() for t in re.split(r"[,/|;]", tech_match.group(1)) if t.strip()]
		else:
			tech_tokens = re.search(r"Tech\s*[:\-]\s*([^\n]+)", line, flags=re.IGNORECASE)
			if tech_tokens:
				tech_stack = [t.strip() for t in re.split(r"[,/|;]", tech_tokens.group(1)) if t.strip()]

		domain = ''
		if link:
			# domain extraction
			try:
				import urllib.parse as _up
				domain = _up.urlparse(link).netloc
			except Exception:
				domain = ''

		# Filter noise: title should be short
		if title and len(title) < 200:
			projects.append({
				'title': title,
				'link': link,
				'description': description,
				'tech_stack': tech_stack,
				'domain': domain,
			})
	return projects


def call_gemini_for_resume(parsed_text: str) -> Optional[Dict[str, Any]]:
	api_key = getattr(settings, "GEMINI_API_KEY", None) or os.getenv("GEMINI_API_KEY")
	if not api_key:
		logger.info("GEMINI_API_KEY not configured; skipping AI extraction")
		return None
	try:
		import google.generativeai as genai
		genai.configure(api_key=api_key)
		model_name = getattr(settings, "GEMINI_MODEL", None) or os.getenv("GEMINI_MODEL") or "gemini-1.5-flash"
		model = genai.GenerativeModel(model_name)
		schema = {
			"type": "object",
			"properties": {
				"is_resume": {"type": "boolean"},
				"name": {"type": "string"},
				"email": {"type": "string"},
				"phone": {"type": "string"},
				"location": {"type": "string"},
				"years_experience": {"type": "number"},
				"skills": {
					"type": "array",
					"items": {
						"type": "object",
						"properties": {
							"name": {"type": "string"},
							"category": {"type": "string"},
							"proficiency": {"type": "string"}
						},
						"required": ["name", "category"]
					}
				},
				"preferred_roles": {"type": "array", "items": {"type": "string"}},
				"job_titles": {"type": "array", "items": {"type": "string"}},
				"suited_roles": {
					"type": "array",
					"items": {
						"type": "object",
						"properties": {
							"role": {"type": "string"},
							"match_percent": {"type": "number"}
						}
					}
				},
				"experiences": {
					"type": "array",
					"items": {
						"type": "object",
						"properties": {
							"title": {"type": "string"},
							"company": {"type": "string"},
							"current": {"type": "boolean"},
							"start_date": {"type": "string"},
							"end_date": {"type": "string"},
							"location": {"type": "string"},
							"description": {"type": "string"}
						}
					}
				},
				"education": {
					"type": "array",
					"items": {
						"type": "object",
						"properties": {
							"degree": {"type": "string"},
							"current": {"type": "boolean"},
							"start_year": {"type": "string"},
							"end_year": {"type": "string"},
							"location": {"type": "string"},
							"description": {"type": "string"},
							"institution": {"type": "string"},
							"field_of_study": {"type": "string"}
						}
					}
				},
				"projects": {
					"type": "array",
					"items": {
						"type": "object",
						"properties": {
							"title": {"type": "string"},
							"link": {"type": "string"},
							"description": {"type": "string"},
							"tech_stack": {"type": "array", "items": {"type": "string"}},
							"domain": {"type": "string"}
						}
					}
				},
				"links": {
					"type": "array",
					"items": {
						"type": "object",
						"properties": {
							"type": {"type": "string"},
							"url": {"type": "string"}
						}
					}
				},
				"linkedin_url": {"type": "string"},
				"portfolio_url": {"type": "string"},
				"github_url": {"type": "string"},
				"website_url": {"type": "string"},
				"expected_salary_range": {
					"type": "object",
					"properties": {
						"min": {"type": "number"},
						"max": {"type": "number"},
						"currency": {"type": "string"}
					}
				},
				"summary": {"type": "string"}
			},
			"required": ["is_resume", "suited_roles", "preferred_roles"]
			}
		# To improve model compliance, include the JSON schema explicitly in the prompt body so the model can follow it.
		schema_text = json.dumps(schema, indent=2)
		prompt = (
			"You are an ATS parser. Extract structured data from this resume text.\n"
			"Return STRICT JSON matching the schema exactly. DO NOT include any explanations, commentary, or surrounding text—only the JSON object.\n\n"
			"SCHEMA:\n" + schema_text + "\n\n"
			"IMPORTANT: produce the exact shapes for 'education' and 'experiences' as described in the schema.\n"
			"For 'skills' return an array of objects with keys: name (string), category (string), proficiency (string). 'category' is REQUIRED for each skill. If the resume does not list a category, infer a reasonable category (e.g., 'Programming', 'Framework', 'Tools', 'Management') rather than leaving it blank.\n"
			"For 'suited_roles' return an array of objects with keys {role (string), match_percent (number)}. Return many candidate roles (up to 20) ordered by match_percent desc. match_percent must be a number between 33 and 100 (inclusive) — use higher percentages for roles that closely match the candidate's skills/experience.\n"
			"After listing suited_roles, compute 'preferred_roles' as those suited roles with match_percent >= 85.0 (use up to top 5 preferred roles); include preferred_roles as a top-level array of role names.\n"
			"CRITICAL: Always include the top-level keys 'suited_roles' and 'preferred_roles' in the JSON output. If none apply, return them as empty arrays. Do NOT omit these keys.\n\n"
			"Return the parsed JSON now. Resume text follows: \n" + parsed_text[:15000]
		)
		resp = model.generate_content(prompt)
		text = resp.text if hasattr(resp, "text") else (resp.candidates[0].content.parts[0].text if getattr(resp, "candidates", None) else "")
		if not text:
			return None
		# Sometimes models wrap in triple backticks
		text = text.strip()
		if text.startswith("```"):
			text = text.strip("`\n ")
			if text.lower().startswith("json"):
				text = text[4:].strip()
		try:
			data = json.loads(text)
		except Exception:
			# Try to repair simple trailing commas or similar issues
			try:
				import re
				cleaned = re.sub(r",\s*([}\]])", r"\1", text)
				data = json.loads(cleaned)
			except Exception as e:
				logger.warning(f"Gemini JSON parse failed: {e}; raw: {text[:200]}…")
				return None
		return data
	except Exception as e:
		logger.error(f"Gemini call failed: {e}")
		return None


def compute_readiness(ai: Optional[Dict[str, Any]], is_resume: Optional[bool] = None) -> Dict[str, Any]:
	"""Compute a readiness score and checklist from parsed AI data.

	If is_resume is explicitly False, returns score 0 and all checks False.
	Scoring weights sum to 100 across key fields.
	"""
	checklist = {
		'name': False,
		'email': False,
		'phone': False,
		'location': False,
		'years_experience': False,
		'skills': False,
		'job_titles': False,
		'education': False,
		'suited_roles': False,
	}
	if ai is None:
		return {'score': 0, 'checklist': checklist}
	if is_resume is False:
		return {'score': 0, 'checklist': checklist}

	checklist['name'] = bool((ai.get('name') or '').strip())
	checklist['email'] = bool((ai.get('email') or '').strip())
	checklist['phone'] = bool((ai.get('phone') or '').strip())
	checklist['location'] = bool((ai.get('location') or '').strip())
	yrs = ai.get('years_experience')
	try:
		yrs_num = int(round(float(yrs))) if yrs is not None else None
	except Exception:
		yrs_num = None
	checklist['years_experience'] = yrs_num is not None
	checklist['skills'] = bool(ai.get('skills'))
	checklist['job_titles'] = bool(ai.get('job_titles'))
	checklist['education'] = bool(ai.get('education'))
	sr = ai.get('suited_roles') or []
	checklist['suited_roles'] = bool(sr)

	weights = {
		'email': 20,
		'phone': 10,
		'name': 10,
		'years_experience': 10,
		'skills': 20,
		'job_titles': 10,
		'education': 10,
		'suited_roles': 5,
		'location': 5,
	}
	score = 0
	for key, weight in weights.items():
		if checklist.get(key):
			score += weight
	score = max(0, min(100, int(score)))
	return {'score': score, 'checklist': checklist}


def merge_into_jobseeker_profile(profile, ai: Dict[str, Any], force: bool = False) -> Tuple[Dict[str, Any], List[str]]:
	"""Merge AI fields into profile.

	By default this is non-destructive and only fills missing fields. If `force=True`,
	AI-provided values will overwrite existing profile fields when present.

	Returns (changes, suited_role_names).

	Rules:
	- Do NOT modify first_name/last_name or email; keep existing user-provided identity.
	- If years_experience missing and no experiences, default to 0.
	- Set preferred_roles from the top suited roles (names only) if not already set.
	- If expected_salary_range present, set salary_min/max and force salary_currency to 'INR'.
	"""
	changes: Dict[str, Any] = {}
	suited: List[str] = []
	if not ai:
		return changes, suited

	def _set(field: str, value: Any):
		"""Set change for field. If force is True, overwrite existing values; otherwise only set when missing."""
		if value in (None, "", [], {}):
			return
		current = getattr(profile, field, None)
		if force or not current:
			changes[field] = value

	# Do NOT update first_name/last_name or email per requirement
	_set("phone", (ai.get("phone") or "").strip())
	_set("location", (ai.get("location") or "").strip())

	# Years of experience; default to 0 if experiences missing
	yrs = ai.get("years_experience")
	experiences = ai.get("experiences") or ai.get("experience") or []
	if yrs is None and not experiences:
		yrs = 0
	if yrs is not None and not getattr(profile, "years_of_experience", None):
		try:
			changes["years_of_experience"] = int(round(float(yrs)))
		except Exception:
			pass

	skills = ai.get("skills") or []
	if skills:
		# Build detailed skills list as objects {name, category, proficiency}
		sk_detailed = []
		for s in skills:
			if isinstance(s, str):
				name = s
				cat = ''
				prof = 'intermediate'
				# if string contains '|' or ':' try to split name|category|proficiency
				if '|' in s or ':' in s:
					parts = re.split(r"[|:]", s)
					if parts:
						name = parts[0].strip()
						if len(parts) > 1:
							cat = parts[1].strip()
						if len(parts) > 2:
							prof = parts[2].strip()
				if name:
					# store as detailed object in skills field per new contract
					sk_detailed.append({"name": name, "category": cat, "proficiency": prof or 'intermediate'})
			elif isinstance(s, dict):
				name = s.get('name') or s.get('skill') or ''
				cat = s.get('category') or s.get('group') or ''
				prof = s.get('proficiency') or s.get('level') or 'intermediate'
				if name:
					sk_detailed.append({"name": name, "category": cat, "proficiency": prof})
		# Write detailed skill objects into the single `skills` field (replace previous skills_detailed)
		if sk_detailed and (force or not getattr(profile, 'skills', None)):
			changes["skills"] = sk_detailed

	edu = ai.get("education") or []
	if edu:
		if isinstance(edu, list) and (force or not getattr(profile, 'education', None)):
			changes["education"] = edu

	projects = ai.get("projects") or []
	if projects and (force or not getattr(profile, 'projects', None)):
		# Normalize project items to expected keys
		normalized = []
		for p in projects:
			if isinstance(p, dict) and (p.get('title') or p.get('description')):
				normalized.append({
					'title': p.get('title') or '',
					'link': p.get('link') or p.get('url') or '',
					'description': p.get('description') or '',
					'tech_stack': p.get('tech_stack') or p.get('techStack') or [],
					'domain': p.get('domain') or ''
				})
			elif isinstance(p, str):
				normalized.append({'title': p, 'link': '', 'description': '', 'tech_stack': [], 'domain': ''})
		if normalized:
			changes['projects'] = normalized

	summary = (ai.get("summary") or "").strip()
	if summary and (force or not getattr(profile, 'bio', None)):
		changes["bio"] = summary

	titles = ai.get("job_titles") or []
	if titles and (force or not getattr(profile, 'job_title', None)):
		if isinstance(titles, list) and titles:
			changes["job_title"] = titles[0]

	# Suited roles may come as strings or objects with {role, match_percent}
	raw_suited = ai.get("suited_roles") or []
	suited_items: List[Dict[str, Any]] = []
	for item in raw_suited:
		if isinstance(item, str):
			suited_items.append({"role": item, "match_percent": None})
		elif isinstance(item, dict) and item.get("role"):
			# normalize percent
			mp = item.get("match_percent")
			try:
				mp = float(mp) if mp is not None else None
			except Exception:
				mp = None
			suited_items.append({"role": item.get("role"), "match_percent": mp})

	# Normalize/clean match_percent values and produce suited objects list
	suited_objects: List[Dict[str, Any]] = []
	for it in suited_items:
		role = (it.get("role") or "").strip()
		if not role:
			continue
		mp = it.get("match_percent")
		try:
			mp_num = float(mp) if mp is not None else None
		except Exception:
			mp_num = None
		# If missing, set a neutral suggested percentage (50)
		if mp_num is None:
			mp_num = 50.0
		# Clamp to requested range 33-100
		mp_num = max(33.0, min(100.0, mp_num))
		suited_objects.append({"role": role, "match_percent": round(mp_num, 2)})

	# Sort suited_objects by match_percent desc and cap candidates to 20
	if suited_objects:
		suited_objects = sorted(suited_objects, key=lambda x: x.get('match_percent', 0), reverse=True)[:20]
	# names list preserved in order
	suited = [o["role"] for o in suited_objects]

	# Preferred roles: compute from suited_objects first (>=85). If none, fall back to AI-provided preferred_roles
	preferred_from_suited = [o for o in suited_objects if o.get("match_percent", 0) >= 85.0]
	if preferred_from_suited and (force or not (getattr(profile, "preferred_roles", None) or [])):
		changes["preferred_roles"] = [r["role"] for r in preferred_from_suited[:5]]
	else:
		# If AI directly provides preferred_roles and we didn't compute any, respect that (force or missing)
		ai_pref = ai.get("preferred_roles") or []
		if ai_pref and (force or not (getattr(profile, "preferred_roles", None) or [])):
			try:
				pr = [str(x).strip() for x in ai_pref if x]
				if pr:
					changes["preferred_roles"] = pr[:5]
			except Exception:
				pass

	# Store suited_job_roles_detailed as list of {role, match_percent}
	if suited_objects and (force or not getattr(profile, 'suited_job_roles_detailed', None)):
		changes['suited_job_roles_detailed'] = suited_objects
	# Do NOT persist legacy `suited_job_roles` list separately; consumers should derive names from detailed field.

	# Expected salary range: set salary_min/max and force INR currency
	expected_salary = ai.get("expected_salary_range") or {}
	if expected_salary and (force or not profile.salary_min or not profile.salary_max or not profile.salary_currency):
		try:
			from decimal import Decimal
			min_val = expected_salary.get("min")
			max_val = expected_salary.get("max")
			if min_val is not None and max_val is not None:
				changes["salary_min"] = Decimal(str(min_val))
				changes["salary_max"] = Decimal(str(max_val))
				changes["salary_currency"] = "INR"
			else:
				# If only one bound, still set currency
				changes.setdefault("salary_currency", "INR")
		except Exception:
			# Ensure currency set to INR if missing
			if not getattr(profile, "salary_currency", None):
				changes["salary_currency"] = "INR"

	# Persist professional links if present and profile fields are empty
	# ai may include explicit fields (linkedin_url, etc.) or a 'links' list; detected_links may be passed separately by caller
	link_sources = {}
	# Prefer ai explicit fields
	for k in ("linkedin_url", "portfolio_url", "github_url", "website_url"):
		v = ai.get(k)
		if v:
			link_sources[k] = v

	# Also normalize from ai['links'] array if present
	for link_obj in ai.get("links", []) or []:
		if not isinstance(link_obj, (dict, str)):
			continue
		if isinstance(link_obj, str):
			url = link_obj
			type_ = ""
		else:
			url = link_obj.get("url") or link_obj.get("link") or ""
			type_ = (link_obj.get("type") or "").lower()
		if not url:
			continue
		u = url.strip()
		if "linkedin" in type_ or "linkedin" in u:
			link_sources.setdefault("linkedin_url", u)
		elif "github" in type_ or "github" in u:
			link_sources.setdefault("github_url", u)
		elif "portfolio" in type_ or "portfolio" in u or "behance" in u or "dribbble" in u:
			link_sources.setdefault("portfolio_url", u)
		else:
			# fallback to website
			link_sources.setdefault("website_url", u)

	# Finally, write into profile only when fields are empty
	try:
		# When force=True we overwrite link fields; otherwise only set when empty
		if (force or getattr(profile, "linkedin_url", None) in (None, "")) and link_sources.get("linkedin_url"):
			changes["linkedin_url"] = link_sources.get("linkedin_url")
		if (force or getattr(profile, "portfolio_url", None) in (None, "")) and link_sources.get("portfolio_url"):
			changes["portfolio_url"] = link_sources.get("portfolio_url")
		if (force or getattr(profile, "github_url", None) in (None, "")) and link_sources.get("github_url"):
			changes["github_url"] = link_sources.get("github_url")
		if (force or getattr(profile, "website_url", None) in (None, "")) and link_sources.get("website_url"):
			changes["website_url"] = link_sources.get("website_url")
	except Exception:
		# Non-fatal: ignore if profile doesn't have these fields
		pass

	return changes, suited

